import os
from motor.motor_asyncio import AsyncIOMotorClient
from datetime import datetime
import uuid
from typing import Dict, List, Optional, Set
from utils.common_utils import to_init_caps
from parsing.parsing_utils import extract_email
from dotenv import load_dotenv
from pathlib import Path
from bson import Binary
from docx import Document
import PyPDF2
import io
import asyncio

# --------------------
# Database Initialization (Async)
# --------------------
async def initialize_mongodb():
    # Load .env from the project root
    env_path = Path(__file__).parent.parent / '.env'
    load_dotenv(dotenv_path=env_path)
    
    MONGO_URI = os.getenv("MONGO_URI")
    MONGO_DB = os.getenv("MONGO_DB")
    
    if not MONGO_URI:
        raise ValueError("MONGO_URI environment variable is not set")
    
    try:
        client = AsyncIOMotorClient(MONGO_URI)
        return client[MONGO_DB]
    except Exception as e:
        raise Exception(f"Failed to initialize MongoDB client: {str(e)}")

# Global db instance
db = None

async def get_db():
    global db
    if db is None:
        db = await initialize_mongodb()
    return db

# --------------------
# Utility Functions (Keep sync for file processing)
# --------------------
import os
import io
import tempfile
import subprocess
import platform
from typing import Optional

import PyPDF2

try:
    import win32com.client  # For Windows MS Word
except ImportError:
    win32com = None

try:
    from docx import Document
except ImportError:
    Document = None

def count_pages(file_content: bytes, filename: str) -> int:
    """
    Count pages in PDF, DOC, or DOCX files.
    Uses:
        - PDF: direct count
        - DOC/DOCX: exact count via LibreOffice (cross-platform) or MS Word COM (Windows)
        - Fallback: paragraph-based estimate
    """
    file_extension = os.path.splitext(filename)[1].lower()

    try:
        # ---------- PDF ----------
        if file_extension == ".pdf":
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            return len(pdf_reader.pages)

        # ---------- DOC/DOCX ----------
        elif file_extension in [".doc", ".docx"]:
            # --- Windows with MS Word ---
            if platform.system() == "Windows" and win32com:
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                    tmp_file.write(file_content)
                    tmp_file_path = tmp_file.name

                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(tmp_file_path)
                pages = doc.ComputeStatistics(2)  # 2 = wdStatisticPages
                doc.Close(False)
                word.Quit()
                os.unlink(tmp_file_path)
                return pages

            # --- LibreOffice cross-platform ---
            else:
                with tempfile.TemporaryDirectory() as tmpdir:
                    input_path = os.path.join(tmpdir, filename)
                    output_pdf = os.path.join(tmpdir, "converted.pdf")

                    with open(input_path, "wb") as f:
                        f.write(file_content)

                    # Convert DOC/DOCX → PDF using LibreOffice
                    subprocess.run(
                        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, input_path],
                        check=True,
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL
                    )

                    if os.path.exists(output_pdf):
                        with open(output_pdf, "rb") as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            return len(pdf_reader.pages)
                    else:
                        # Fallback if conversion failed
                        if Document:
                            doc = Document(io.BytesIO(file_content))
                            return max(1, len(doc.paragraphs) // 10)
                        return 1

        # ---------- Other file types ----------
        else:
            return 1

    except Exception as e:
        print(f"Error counting pages for {filename}: {e}")
        # Fallback to paragraph-based estimate for DOCX
        if Document and file_extension in [".doc", ".docx"]:
            try:
                doc = Document(io.BytesIO(file_content))
                return max(1, len(doc.paragraphs) // 10)
            except:
                return 1
        return 1

# --------------------
# Client Management Functions (Async)
# --------------------
async def fetch_client_names(company_id: str, status: Optional[str] = None) -> Set[str]:
    try:
        db = await get_db()
        query = {"company_id": company_id, "is_deleted": False}
        if status:
            query["status"] = status
            
        client_names = await db.clients.distinct("client_name", query)
        return set(sorted(client_names))
    except Exception as e:
        raise Exception(f"Failed to fetch client names: {str(e)}")

async def fetch_client_details(client_name: str, company_id: str) -> Optional[Dict]:
    try:
        db = await get_db()
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id,
            "is_deleted": False
        })
        
        if not client_doc:
            return None
            
        client_id = client_doc["_id"]
        
        # Get the most recent job description for this client
        jd_doc = await db.job_descriptions.find_one(
            {"client_id": client_id, "company_id": company_id, "is_deleted": False},
            sort=[("created_at", -1)]
        )
        
        if not jd_doc:
            return None
            
        return {
            "job_description": jd_doc.get("jd_title", ""),
            "required_experience": jd_doc.get("required_experience", ""),
            "primary_skills": jd_doc.get("primary_skills", []),
            "secondary_skills": jd_doc.get("secondary_skills", []),
            "status": jd_doc.get("status", "active")
        }
    except Exception as e:
        raise Exception(f"Failed to fetch client details: {str(e)}")

async def fetch_client_details_by_jd(client_name: str, jd_name: str, company_id: str) -> Optional[Dict]:
    try:
        db = await get_db()
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id,
            "is_deleted": False
        })
        
        if not client_doc:
            return None
            
        client_id = client_doc["_id"]
        
        jd_doc = await db.job_descriptions.find_one({
            "client_id": client_id,
            "jd_title": await to_init_caps(jd_name),
            "company_id": company_id,
            "is_deleted": False
        })
        
        if not jd_doc:
            return None
            
        return {
            "job_description": jd_doc.get("jd_title", ""),
            "required_experience": jd_doc.get("required_experience", ""),
            "primary_skills": jd_doc.get("primary_skills", []),
            "secondary_skills": jd_doc.get("secondary_skills", []),
            "status": jd_doc.get("status", "active")
        }
    except Exception as e:
        raise Exception(f"Failed to fetch client details by JD: {str(e)}")

async def fetch_jd_names_for_client(client_name: str, company_id: str, status: Optional[str] = None) -> Optional[List[str]]:
    try:
        db = await get_db()
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id,
            "is_deleted": False
        })
        
        if not client_doc:
            return None
            
        client_id = client_doc["_id"]
        
        query = {"client_id": client_id, "company_id": company_id, "is_deleted": False}
        if status:
            query["status"] = status
            
        jd_names = await db.job_descriptions.distinct("jd_title", query)
        
        return jd_names if jd_names else None
    except Exception as e:
        raise Exception(f"Failed to fetch JD names for client: {str(e)}")

async def update_client_status(client_name: str, company_id: str, status: str) -> bool:
    """
    Update the status of a client
    """
    try:
        if status not in ["active", "inactive"]:
            raise ValueError("Status must be 'active' or 'inactive'")
        
        db = await get_db()
        result = await db.clients.update_one(
            {
                "client_name": await to_init_caps(client_name),
                "company_id": company_id,
                "is_deleted": False
            },
            {"$set": {"status": status}}
        )
        
        return result.modified_count > 0
    except Exception as e:
        print(f"Failed to update client status: {e}")
        return False

async def soft_delete_client(client_name: str, company_id: str, deleted_by: str) -> bool:
    """
    Soft delete a client and all its associated JDs and analyses
    """
    try:
        db = await get_db()
        # Soft delete the client
        client_result = await db.clients.update_one(
            {
                "client_name": await to_init_caps(client_name),
                "company_id": company_id
            },
            {
                "$set": {
                    "is_deleted": True,
                    "status": "inactive",
                    "deleted_at": datetime.now(),
                    "deleted_by": deleted_by
                }
            }
        )
        
        if client_result.modified_count == 0:
            return False
            
        # Get client ID
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id
        })
        
        if not client_doc:
            return False
            
        client_id = client_doc["_id"]
        
        # Soft delete all JDs for this client
        await db.job_descriptions.update_many(
            {
                "client_id": client_id,
                "company_id": company_id
            },
            {
                "$set": {
                    "is_deleted": True,
                    "status": "inactive",
                    "deleted_at": datetime.now(),
                    "deleted_by": deleted_by
                }
            }
        )
        
        # Soft delete all analyses for this client
        await db.analysis_history.update_many(
            {
                "client_id": client_id,
                "company_id": company_id
            },
            {
                "$set": {
                    "is_deleted": True,
                    "status": "inactive",
                    "deleted_at": datetime.now(),
                    "deleted_by": deleted_by
                }
            }
        )
        
        return True
    except Exception as e:
        print(f"Failed to soft delete client: {e}")
        return False

async def restore_client(client_name: str, company_id: str, restored_by: str) -> bool:
    """
    Restore a soft-deleted client and its associated JDs and analyses
    """
    try:
        db = await get_db()
        # Restore the client
        client_result = await db.clients.update_one(
            {
                "client_name": await to_init_caps(client_name),
                "company_id": company_id,
                "is_deleted": True
            },
            {
                "$set": {
                    "is_deleted": False,
                    "status": "active",
                    "restored_at": datetime.now(),
                    "restored_by": restored_by
                },
                "$unset": {
                    "deleted_at": "",
                    "deleted_by": ""
                }
            }
        )
        
        if client_result.modified_count == 0:
            return False
            
        # Get client ID
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id
        })
        
        if not client_doc:
            return False
            
        client_id = client_doc["_id"]
        
        # Restore all JDs for this client
        await db.job_descriptions.update_many(
            {
                "client_id": client_id,
                "company_id": company_id,
                "is_deleted": True
            },
            {
                "$set": {
                    "is_deleted": False,
                    "status": "active",
                    "restored_at": datetime.now(),
                    "restored_by": restored_by
                },
                "$unset": {
                    "deleted_at": "",
                    "deleted_by": ""
                }
            }
        )
        
        # Restore all analyses for this client
        await db.analysis_history.update_many(
            {
                "client_id": client_id,
                "company_id": company_id,
                "is_deleted": True
            },
            {
                "$set": {
                    "is_deleted": False,
                    "status": "active",
                    "restored_at": datetime.now(),
                    "restored_by": restored_by
                },
                "$unset": {
                    "deleted_at": "",
                    "deleted_by": ""
                }
            }
        )
        
        return True
    except Exception as e:
        print(f"Failed to restore client: {e}")
        return False

# --------------------
# Job Description Management Functions (Async)
# --------------------
async def update_job_description(client_name: str, jd_name: str, required_experience: str, 
                               primary_skills: list, secondary_skills: list, company_id: str) -> bool:
    """
    Update the job description details for a given client and JD name in MongoDB
    """
    try:
        db = await get_db()
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id,
            "is_deleted": False
        })
        if not client_doc:
            return False
            
        client_id = client_doc["_id"]
        
        jd_doc = await db.job_descriptions.find_one({
            "client_id": client_id,
            "jd_title": await to_init_caps(jd_name),
            "company_id": company_id,
            "is_deleted": False
        })
        if not jd_doc:
            return False
            
        result = await db.job_descriptions.update_one(
            {"_id": jd_doc["_id"]},
            {"$set": {
                "required_experience": required_experience,
                "primary_skills": primary_skills,
                "secondary_skills": secondary_skills
            }}
        )
        
        return result.modified_count > 0
    except Exception as e:
        print(f"Failed to update job description: {e}")
        return False

async def update_jd_status(client_name: str, jd_title: str, company_id: str, status: str) -> bool:
    """
    Update the status of a job description
    """
    try:
        if status not in ["active", "inactive"]:
            raise ValueError("Status must be 'active' or 'inactive'")
        
        db = await get_db()
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id,
            "is_deleted": False
        })
        
        if not client_doc:
            return False
            
        client_id = client_doc["_id"]
        
        result = await db.job_descriptions.update_one(
            {
                "client_id": client_id,
                "jd_title": await to_init_caps(jd_title),
                "company_id": company_id,
                "is_deleted": False
            },
            {"$set": {"status": status}}
        )
        
        return result.modified_count > 0
    except Exception as e:
        print(f"Failed to update JD status: {e}")
        return False

async def soft_delete_jd(client_name: str, jd_title: str, company_id: str, deleted_by: str) -> bool:
    """
    Soft delete a job description and all its associated analyses
    """
    try:
        db = await get_db()
        # Get client ID
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id
        })
        
        if not client_doc:
            return False
            
        client_id = client_doc["_id"]
        
        # Soft delete the JD
        jd_result = await db.job_descriptions.update_one(
            {
                "client_id": client_id,
                "jd_title": await to_init_caps(jd_title),
                "company_id": company_id
            },
            {
                "$set": {
                    "is_deleted": True,
                    "status": "inactive",
                    "deleted_at": datetime.now(),
                    "deleted_by": deleted_by
                }
            }
        )
        
        if jd_result.modified_count == 0:
            return False
            
        # Soft delete all analyses for this JD
        await db.analysis_history.update_many(
            {
                "jd_id": client_id,
                "company_id": company_id
            },
            {
                "$set": {
                    "is_deleted": True,
                    "status": "inactive",
                    "deleted_at": datetime.now(),
                    "deleted_by": deleted_by
                }
            }
        )
        
        return True
    except Exception as e:
        print(f"Failed to soft delete JD: {e}")
        return False

async def restore_jd(client_name: str, jd_title: str, company_id: str, restored_by: str) -> bool:
    """
    Restore a soft-deleted job description and its associated analyses
    """
    try:
        db = await get_db()
        # Get client ID
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id
        })
        
        if not client_doc:
            return False
            
        client_id = client_doc["_id"]
        
        # Restore the JD
        jd_result = await db.job_descriptions.update_one(
            {
                "client_id": client_id,
                "jd_title": await to_init_caps(jd_title),
                "company_id": company_id,
                "is_deleted": True
            },
            {
                "$set": {
                    "is_deleted": False,
                    "status": "active",
                    "restored_at": datetime.now(),
                    "restored_by": restored_by
                },
                "$unset": {
                    "deleted_at": "",
                    "deleted_by": ""
                }
            }
        )
        
        if jd_result.modified_count == 0:
            return False
            
        # Restore all analyses for this JD
        await db.analysis_history.update_many(
            {
                "jd_id": client_id,
                "company_id": company_id,
                "is_deleted": True
            },
            {
                "$set": {
                    "is_deleted": False,
                    "status": "active",
                    "restored_at": datetime.now(),
                    "restored_by": restored_by
                },
                "$unset": {
                    "deleted_at": "",
                    "deleted_by": ""
                }
            }
        )
        
        return True
    except Exception as e:
        print(f"Failed to restore JD: {e}")
        return False

# --------------------
# Analysis Management Functions
# --------------------
async def store_results_in_mongodb(
    analysis_data: Dict,
    jd_data: Dict,
    filename: str,
    resume_text: str,
    file_content: bytes,
    client_name: str,
    job_description: str,
    created_by: str,
    company_id: str,
    name: str,
    role: str
) -> Optional[str]:
    try:
        db = await get_db()
        page_count = count_pages(file_content, filename)
        current_time = datetime.now()

        # ===== CLIENT SECTION =====
        client_doc = await db.clients.find_one({
            "client_name": await to_init_caps(client_name),
            "company_id": company_id,
            "is_deleted": False
        })

        if client_doc:
            client_id = client_doc["_id"]
        else:
            new_client = {
                "client_name": await to_init_caps(client_name),
                "company_id": company_id,
                "created_by": created_by,
                "created_at": current_time,
                "status": "active",
                "is_deleted": False
            }
            result = await db.clients.insert_one(new_client)
            client_id = result.inserted_id
            await log_audit_trail(
                user_id=created_by,
                name=name,
                role=role,
                company_id=company_id,
                action="create_client",
                target_table="clients",
                target_id=str(client_id),
                old_data={},
                new_data={"client_name": new_client["client_name"]},
                screen="users"
            )
            client_doc = new_client

        # ===== JOB DESCRIPTION SECTION =====
        jd_doc = await db.job_descriptions.find_one({
            "client_id": client_id,
            "jd_title": await to_init_caps(job_description),
            "company_id": company_id,
            "is_deleted": False
        })

        if jd_doc:
            jd_id = jd_doc["_id"]
        else:
            new_jd = {
                "client_id": client_id,
                "jd_title": await to_init_caps(job_description),
                "required_experience": jd_data.get("required_experience", ""),
                "primary_skills": jd_data.get("primary_skills", []),
                "secondary_skills": jd_data.get("secondary_skills", []),
                "location": jd_data.get("location"),
                "budget": jd_data.get("budget"),
                "number_of_positions": jd_data.get("number_of_positions"),
                "work_mode": jd_data.get("work_mode"),
                "company_id": company_id,
                "created_by": created_by,
                "created_at": current_time,
                "status": "active",
                "is_deleted": False
            }
            result = await db.job_descriptions.insert_one(new_jd)
            jd_id = result.inserted_id
            await log_audit_trail(
                user_id=created_by,
                name=name,
                role=role,
                company_id=company_id,
                action="create_job_description",
                target_table="job_descriptions",
                target_id=str(jd_id),
                old_data={},
                new_data={
                    "jd_title": new_jd["jd_title"],
                    "required_experience": new_jd.get("required_experience", ""),
                    "primary_skills": new_jd.get("primary_skills", []),
                    "secondary_skills": new_jd.get("secondary_skills", []),
                    "location": new_jd.get("location"),
                    "budget": new_jd.get("budget")
                },
                screen="users"
            )
            jd_doc = new_jd

        # ===== ANALYSIS SECTION =====
        analysis_id = str(uuid.uuid4())
        candidate_email = extract_email(resume_text)
        candidate_name = analysis_data.get("candidate_info", {}).get("candidate_name", "Not specified")
        profile_feedback = analysis_data.get("profile_feedback", {})

        analysis_record = {
            "analysis_id": analysis_id,
            "timestamp": current_time,
            "candidate_name": candidate_name,
            "filename": filename,
            "file_content": Binary(file_content),
            "page_count": page_count,
            "client_id": client_id,
            "client_name": client_doc["client_name"],
            "jd_id": jd_doc["_id"],
            "jd_title": jd_doc["jd_title"],
            "required_experience": jd_doc.get("required_experience", ""),
            "primary_skills": jd_doc.get("primary_skills", []),
            "secondary_skills": jd_doc.get("secondary_skills", []),
            "candidate_email": candidate_email,
            "freelancer_status": profile_feedback.get("freelancer_status", False),
            "has_linkedin": profile_feedback.get("has_linkedin", False),
            "linkedin_url": profile_feedback.get("linkedin_url", ""),
            "has_email": profile_feedback.get("has_email", False),
            "match_score": analysis_data.get("skill_analysis", {}).get("match_score", 0),
            "experience_match": analysis_data.get("experience_analysis", {}).get("experience_match", False),
            "total_experience": analysis_data.get("experience_analysis", {}).get("total_experience", "N/A"),
            "matching_skills": analysis_data.get("skill_analysis", {}).get("matching_skills", []),
            "missing_primary_skills": analysis_data.get("skill_analysis", {}).get("missing_primary_skills", []),
            "missing_secondary_skills": analysis_data.get("skill_analysis", {}).get("missing_secondary_skills", []),
            "company_id": company_id,
            "created_by": created_by,
            "status": "active",
            "is_deleted": False,
        }

        await db.analysis_history.insert_one(analysis_record)

        return analysis_id

    except Exception as e:
        raise Exception(f"Failed to store results in MongoDB: {str(e)}")

async def fetch_analysis_history(current_user: dict) -> List[Dict]:
    try:
        db = await get_db()
        # Build query based on user role - only for analysis history
        if current_user["role"] == "company_admin":
            # Company admin can see all analyses for their company
            query = {"company_id": current_user["company_id"],
                     "is_deleted": False}
        elif current_user["role"] == "user":
            # Regular user can only see their own analyses within their company
            query = {
                "company_id": current_user["company_id"],
                "created_by": current_user["user_id"],
                "is_deleted": False
            }
        else:
            # For other roles, return empty
            query = {"company_id": "invalid_id"}  # Ensure no results
            
        # Exclude file_content from the query to reduce payload size
        history_cursor = db.analysis_history.find(
            query, 
            {"file_content": 0}  # Exclude file content from results
        ).sort("timestamp", -1)
        
        history = await history_cursor.to_list(length=None)
        
        # Convert ObjectId to string for JSON serialization
        for item in history:
            item["_id"] = str(item["_id"])
            if "client_id" in item:
                item["client_id"] = str(item["client_id"])
            if "jd_id" in item:
                item["jd_id"] = str(item["jd_id"])
        return history
    
    except Exception as e:
        raise Exception(f"Failed to fetch history: {str(e)}")

async def update_analysis_status(analysis_id: str, company_id: str, status: str) -> bool:
    """
    Update the status of an analysis
    """
    try:
        if status not in ["active", "inactive"]:
            raise ValueError("Status must be 'active' or 'inactive'")
        
        db = await get_db()
        result = await db.analysis_history.update_one(
            {
                "analysis_id": analysis_id,
                "company_id": company_id,
                "is_deleted": False
            },
            {"$set": {"status": status}}
        )
        
        return result.modified_count > 0
    except Exception as e:
        print(f"Failed to update analysis status: {e}")
        return False

async def soft_delete_analysis(analysis_id: str, company_id: str, deleted_by: str) -> bool:
    """
    Soft delete an analysis by setting is_deleted to True and status to inactive
    """
    try:
        db = await get_db()
        result = await db.analysis_history.update_one(
            {
                "analysis_id": analysis_id,
                "company_id": company_id
            },
            {
                "$set": {
                    "is_deleted": True,
                    "status": "inactive",
                    "deleted_at": datetime.now(),
                    "deleted_by": deleted_by
                }
            }
        )
        return result.modified_count > 0
    except Exception as e:
        print(f"Failed to soft delete analysis: {e}")
        return False

async def restore_analysis(analysis_id: str, company_id: str, restored_by: str) -> bool:
    """
    Restore a soft-deleted analysis
    """
    try:
        db = await get_db()
        result = await db.analysis_history.update_one(
            {
                "analysis_id": analysis_id,
                "company_id": company_id,
                "is_deleted": True
            },
            {
                "$set": {
                    "is_deleted": False,
                    "status": "active",
                    "restored_at": datetime.now(),
                    "restored_by": restored_by
                },
                "$unset": {
                    "deleted_at": "",
                    "deleted_by": ""
                }
            }
        )
        return result.modified_count > 0
    except Exception as e:
        print(f"Failed to restore analysis: {e}")
        return False

# --------------------
# Usage Tracking Functions (Async)
# --------------------
async def initialize_usage_tracking(company_id: str):
    """Initialize usage tracking for a new company"""
    db = await get_db()
    current_month = datetime.now().strftime("%Y-%m")
    usage_record = {
        "company_id": company_id,
        "month": current_month,
        "page_count": 0,
        "last_updated": datetime.now()
    }
    await db.company_usage.insert_one(usage_record)

async def get_current_month_usage(company_id: str) -> int:
    """Get current month's page usage for a company"""
    db = await get_db()
    current_month = datetime.now().strftime("%Y-%m")
    usage_record = await db.company_usage.find_one({
        "company_id": company_id,
        "month": current_month
    })
    return usage_record["page_count"] if usage_record else 0

async def increment_usage(company_id: str, page_count: int):
    """Increment company's page usage"""
    db = await get_db()
    current_month = datetime.now().strftime("%Y-%m")
    
    result = await db.company_usage.update_one(
        {
            "company_id": company_id,
            "month": current_month
        },
        {
            "$inc": {"page_count": page_count},
            "$set": {"last_updated": datetime.now()}
        },
        upsert=True
    )
    
    return result.modified_count > 0

async def get_company_page_limit(company_id: str) -> int:
    """Get company's monthly page limit"""
    db = await get_db()
    company = await db.companies.find_one(
        {"id": company_id, "is_deleted": False},
        {"monthly_page_limit": 1}
    )
    return company.get("monthly_page_limit", 1000) if company else 1000

async def check_usage_limit(company_id: str, additional_pages: int = 0) -> bool:
    """Check if company has reached its monthly limit"""
    current_usage = await get_current_month_usage(company_id)
    page_limit = await get_company_page_limit(company_id)
    
    return (current_usage + additional_pages) <= page_limit

async def get_usage_stats(company_id: str) -> Dict:
    """Get comprehensive usage statistics"""
    current_month = datetime.now().strftime("%Y-%m")
    current_usage = await get_current_month_usage(company_id)
    page_limit = await get_company_page_limit(company_id)
    
    return {
        "current_month": current_month,
        "current_usage": current_usage,
        "page_limit": page_limit,
        "remaining_pages": max(0, page_limit - current_usage),
        "usage_percentage": (current_usage / page_limit * 100) if page_limit > 0 else 0
    }

# --------------------
# Deleted Items Functions (Async)
# --------------------
async def get_deleted_clients(company_id: str):
    db = await get_db()
    clients_cursor = db.clients.find({
        "company_id": company_id,
        "is_deleted": True
    }, {"_id": 0})
    return await clients_cursor.to_list(length=None)

async def get_deleted_jds(company_id: str):
    db = await get_db()
    jds_cursor = db.job_descriptions.find({
        "company_id": company_id,
        "is_deleted": True
    }, {"_id": 0})
    return await jds_cursor.to_list(length=None)

async def get_deleted_analyses(company_id: str):
    db = await get_db()
    analyses_cursor = db.analysis_history.find({
        "company_id": company_id,
        "is_deleted": True
    }, {"file_content": 0})
    return await analyses_cursor.to_list(length=None)

# --------------------
# Status Filtering Functions (Async)
# --------------------
async def get_clients_by_status(company_id: str, status: str):
    db = await get_db()
    clients_cursor = db.clients.find({
        "company_id": company_id,
        "is_deleted": False,
        "status": status
    }, {"_id": 0})
    return await clients_cursor.to_list(length=None)

async def get_jds_by_status(company_id: str, status: str):
    db = await get_db()
    jds_cursor = db.job_descriptions.find({
        "company_id": company_id,
        "is_deleted": False,
        "status": status
    }, {"_id": 0})
    return await jds_cursor.to_list(length=None)

async def get_analyses_by_status(company_id: str, status: str):
    db = await get_db()
    analyses_cursor = db.analysis_history.find({
        "company_id": company_id,
        "is_deleted": False,
        "status": status
    }, {"file_content": 0})
    return await analyses_cursor.to_list(length=None)

# --------------------
# Audit Log Functions (Async)
# --------------------
async def log_audit_trail(
    user_id: str = None, 
    name: str = None,
    role: str = None,
    company_id: str = None,  
    action: str = None, 
    target_table: str = None, 
    target_id: str = None,
    old_data: dict = None,
    new_data: dict = None,
    screen: str = None 
):
    db = await get_db()
    log = {
        "user_id": user_id,
        "name": name,
        "role": role,
        "company_id": company_id,
        "action": action,
        "target_table": target_table,
        "target_id": target_id,
        "old_data": old_data,
        "new_data": new_data,
        "screen": screen,
        "timestamp": datetime.utcnow()
    }
    await db.audit_logs.insert_one(log)
